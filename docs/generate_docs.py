"""
Generate Word Documents for Canvex Documentation
Creates USER_GUIDE.docx and TECHNICAL_DOCS.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Colors
ACCENT_BLUE = RGBColor(10, 132, 255)
DARK_TEXT = RGBColor(29, 29, 31)
GRAY_TEXT = RGBColor(110, 110, 115)


def set_cell_shading(cell, color):
    """Set background color for table cell"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = ACCENT_BLUE
            run.font.size = Pt(24)
        elif level == 2:
            run.font.color.rgb = ACCENT_BLUE
            run.font.size = Pt(18)
        else:
            run.font.color.rgb = DARK_TEXT
            run.font.size = Pt(14)
    return heading


def add_table(doc, headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, 'E8E8ED')
        for paragraph in cell.paragraphs:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.size = Pt(10)
    
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    return table


def create_user_guide():
    """Create the User Guide Word document"""
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Canvex User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(36)
    
    subtitle = doc.add_paragraph('🖼️ Image Excel Creator')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = GRAY_TEXT
    
    tagline = doc.add_paragraph('Automatically search and insert images into Excel files')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].italic = True
    
    version = doc.add_paragraph('Version 1.0 | © 2025 Kunal Pagariya')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.color.rgb = GRAY_TEXT
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Getting Started',
        '3. Main Interface Overview',
        '4. Step-by-Step Workflow',
        '5. Configuration Options',
        '6. Column Mappings',
        '7. Settings Panel',
        '8. Output Files',
        '9. Tips & Best Practices',
        '10. Troubleshooting',
        '11. FAQ'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # Section 1: Introduction
    add_heading(doc, '1. Introduction', 1)
    
    add_heading(doc, 'What is Canvex?', 2)
    doc.add_paragraph(
        'Canvex is a powerful desktop application that automatically searches the web for images '
        'based on text in your Excel spreadsheet and inserts them directly into a new Excel file.'
    )
    
    add_heading(doc, 'Perfect for:', 3)
    uses = doc.add_paragraph()
    uses.add_run('📸 Creating employee directories with headshots\n').bold = False
    uses.add_run('🎬 Building cast lists with actor photos\n')
    uses.add_run('🏢 Generating product catalogs with images\n')
    uses.add_run('📊 Any data visualization requiring images')
    
    add_heading(doc, 'Key Features', 2)
    add_table(doc, 
        ['Feature', 'Description'],
        [
            ['🔍 Multi-Engine Search', 'Search using Bing, Google, or DuckDuckGo'],
            ['🎨 Smart Filtering', 'Automatically removes low-quality, B&W, and cartoon images'],
            ['⚡ Parallel Processing', 'Downloads multiple images simultaneously'],
            ['💾 Auto-Save Settings', 'Your preferences are remembered between sessions'],
            ['🌓 Theme Support', 'Light, Dark, or System-following themes'],
            ['📐 Flexible Resolution', 'From 240p to 4K, or custom values'],
            ['🎯 Portrait Priority', 'Prefers portrait-oriented images for headshots']
        ]
    )
    
    doc.add_page_break()
    
    # Section 2: Getting Started
    add_heading(doc, '2. Getting Started', 1)
    
    add_heading(doc, 'System Requirements', 2)
    add_table(doc,
        ['Requirement', 'Specification'],
        [
            ['Operating System', 'macOS 10.14+ or Windows 10+'],
            ['Internet Connection', 'Required for image searches'],
            ['Chrome Browser', 'Must be installed (used for web scraping)'],
            ['RAM', '4GB minimum, 8GB recommended'],
            ['Storage', '100MB for app + space for output files']
        ]
    )
    
    add_heading(doc, 'Quick Start (5 Minutes)', 2)
    steps = doc.add_paragraph()
    steps.add_run('1. Load Excel').bold = True
    steps.add_run(' → ')
    steps.add_run('2. Set Theme').bold = True
    steps.add_run(' → ')
    steps.add_run('3. Add Mappings').bold = True
    steps.add_run(' → ')
    steps.add_run('4. Start').bold = True
    steps.add_run(' → ')
    steps.add_run('5. Save').bold = True
    
    doc.add_page_break()
    
    # Section 3: Main Interface
    add_heading(doc, '3. Main Interface Overview', 1)
    
    add_heading(doc, 'Toolbar Buttons', 2)
    add_table(doc,
        ['Button', 'Function'],
        [
            ['File', 'Open files, view recent files, reveal settings location'],
            ['Settings', 'Configure search engine, filters, resolution, and format'],
            ['Help', 'View user guide within the app'],
            ['About', 'Application version and contact information'],
            ['Theme', 'Switch between Light, Dark, or System theme']
        ]
    )
    
    doc.add_page_break()
    
    # Section 4: Step-by-Step Workflow
    add_heading(doc, '4. Step-by-Step Workflow', 1)
    
    add_heading(doc, 'Step 1: Load Your Excel File', 2)
    doc.add_paragraph('Option A: Click to Browse')
    p = doc.add_paragraph()
    p.add_run('1. Click the "Select Excel File" button\n')
    p.add_run('2. Navigate to your .xlsx file\n')
    p.add_run('3. Click Open')
    
    doc.add_paragraph()
    doc.add_paragraph('Option B: Drag and Drop')
    p = doc.add_paragraph()
    p.add_run('1. Open your file explorer/finder\n')
    p.add_run('2. Drag the .xlsx file onto the Canvex window\n')
    p.add_run('3. Release to load')
    
    tip = doc.add_paragraph()
    tip.add_run('💡 Tip: ').bold = True
    tip.add_run('Only .xlsx files are supported. Convert older .xls files first.')
    
    add_heading(doc, 'Step 2: Configure Image Settings', 2)
    
    add_heading(doc, 'Image Theme Options', 3)
    add_table(doc,
        ['Theme', 'Best For'],
        [
            ['headshot portrait closeup face', 'Professional headshots, ID photos'],
            ['cinematic lighting portrait', 'Dramatic, artistic portraits'],
            ['studio headshot clean background', 'Corporate/LinkedIn style photos'],
            ['dramatic portrait closeup', 'High-contrast artistic shots'],
            ['smiling closeup face', 'Friendly, approachable photos'],
            ['full body portrait', 'Full-length photos'],
            ['natural daylight portrait', 'Outdoor, natural lighting'],
            ['magazine cover portrait', 'High-fashion style'],
            ['Custom Theme...', 'Enter your own search keywords']
        ]
    )
    
    add_heading(doc, 'Search Browser', 3)
    add_table(doc,
        ['Engine', 'Characteristics'],
        [
            ['Bing Images', '⭐ Recommended. Fastest and most reliable'],
            ['Google Images', 'Alternative results, may be slower'],
            ['DuckDuckGo', 'Privacy-focused, good backup option']
        ]
    )
    
    add_heading(doc, 'Resolution', 3)
    add_table(doc,
        ['Setting', 'Pixels', 'Use Case', 'Speed'],
        [
            ['240p', '240', 'Thumbnails', '●●●● Fastest'],
            ['360p', '360', 'Small previews', '●●●○'],
            ['480p', '480', 'Standard docs', '●●○○'],
            ['720p', '720', '⭐ Recommended', '●●○○'],
            ['1080p', '1080', 'High-quality', '●○○○'],
            ['1440p', '1440', 'Large displays', '○○○○ Slowest'],
            ['2160p', '2160', '4K quality', '○○○○'],
            ['Custom...', '240-4000', 'Your choice', 'Varies']
        ]
    )
    
    add_heading(doc, 'Image Format', 3)
    add_table(doc,
        ['Format', 'Quality', 'File Size', 'Transparency'],
        [
            ['PNG', '★★★ Best', 'Large', '✓ Yes'],
            ['JPG', '★★ Good', 'Medium', '✗ No'],
            ['WEBP', '★★ Good', 'Smallest', '✓ Yes']
        ]
    )
    
    doc.add_page_break()
    
    add_heading(doc, 'Step 3: Set Up Column Mappings', 2)
    doc.add_paragraph(
        'Column mappings tell Canvex which columns contain search terms and where to put the images.'
    )
    
    p = doc.add_paragraph()
    p.add_run('1. Click "').bold = False
    p.add_run('+ Add Mapping').bold = True
    p.add_run('" button\n')
    p.add_run('2. Configure the mapping row:')
    
    add_table(doc,
        ['Field', 'Description', 'Example'],
        [
            ['Input Column', 'Column with search text', 'actor_name'],
            ['Output Column', 'Where to insert images', 'actor_image'],
            ['New Column Name', 'For new columns only', 'photo']
        ]
    )
    
    add_heading(doc, 'Step 4: Start Processing', 2)
    p = doc.add_paragraph()
    p.add_run('1. Click "▶ Start Processing" (green button)\n')
    p.add_run('2. Choose where to save the output file\n')
    p.add_run('3. Enter a filename (e.g., output_with_images.xlsx)\n')
    p.add_run('4. Click Save')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('During Processing:\n').bold = True
    p.add_run('• The progress bar shows overall completion\n')
    p.add_run('• The Cancel button lets you stop safely\n')
    p.add_run('• Processing continues in the background')
    
    add_heading(doc, 'Step 5: Review Output', 2)
    doc.add_paragraph('When complete, a dialog appears asking if you want to open the file.')
    
    add_table(doc,
        ['File', 'Contents'],
        [
            ['your_output.xlsx', 'Excel with images inserted'],
            ['your_output_log.txt', 'Processing log (always created)'],
            ['your_output_ERROR_log.txt', 'Error details (only if errors occurred)']
        ]
    )
    
    doc.add_page_break()
    
    # Section 5: Configuration Options
    add_heading(doc, '5. Configuration Options', 1)
    
    add_heading(doc, 'Image Theme Details', 2)
    doc.add_paragraph('The theme affects search query construction:')
    p = doc.add_paragraph()
    p.add_run('Search Query = [Cell Value] + [Theme]\n').bold = False
    p.add_run('Example: "Tom Hanks" + "headshot portrait closeup face"').italic = True
    
    add_heading(doc, 'Custom Theme Tips', 3)
    p = doc.add_paragraph()
    p.add_run('• Use descriptive words: professional, corporate, natural\n')
    p.add_run('• Add style modifiers: high quality, HD, portrait\n')
    p.add_run('• Specify background: white background, studio')
    
    add_heading(doc, 'Search Browser Comparison', 2)
    add_table(doc,
        ['Feature', 'Bing', 'Google', 'DuckDuckGo'],
        [
            ['Speed', '●●●●', '●●○○', '●●●○'],
            ['Reliability', '●●●●', '●●●○', '●●●○'],
            ['Image Quality', '●●●●', '●●●●', '●●●○'],
            ['Rate Limiting', 'Low', 'Medium', 'Low'],
            ['Fallback', '—', 'Bing', '—']
        ]
    )
    
    note = doc.add_paragraph()
    note.add_run('📝 Note: ').bold = True
    note.add_run('If Google returns no results, Canvex automatically tries Bing as a fallback.')
    
    doc.add_page_break()
    
    # Section 6: Column Mappings
    add_heading(doc, '6. Column Mappings', 1)
    
    add_heading(doc, 'Creating New Columns', 2)
    p = doc.add_paragraph()
    p.add_run('1. In Output Column, select "Create New Column..."\n')
    p.add_run('2. A text field appears\n')
    p.add_run('3. Enter the new column name\n')
    p.add_run('4. The new column is added to the right of existing data')
    
    add_heading(doc, 'Multiple Mappings Example', 2)
    doc.add_paragraph('Input Excel:')
    add_table(doc,
        ['lead_actor', 'supporting_actor', 'director'],
        [
            ['Tom Cruise', 'Val Kilmer', 'Tony Scott'],
            ['Keanu Reeves', 'Laurence Fishburne', 'The Wachowskis']
        ]
    )
    
    doc.add_paragraph('Mappings:')
    p = doc.add_paragraph()
    p.add_run('1. lead_actor → lead_photo (new)\n')
    p.add_run('2. supporting_actor → support_photo (new)\n')
    p.add_run('3. director → director_photo (new)')
    
    doc.add_page_break()
    
    # Section 7: Settings Panel
    add_heading(doc, '7. Settings Panel', 1)
    doc.add_paragraph('Access via Settings button in the toolbar.')
    
    add_heading(doc, 'Image Filters', 2)
    add_table(doc,
        ['Filter', 'Effect', 'Recommended For'],
        [
            ['Prioritize portrait images', 'Prefers taller-than-wide images', 'Headshots, portraits'],
            ['Filter out B&W images', 'Excludes grayscale images', 'Modern, colorful photos'],
            ['Filter out cartoon images', 'Excludes illustrations/graphics', 'Real photographs only']
        ]
    )
    
    add_heading(doc, 'Filter Recommendation Matrix', 3)
    add_table(doc,
        ['Use Case', 'Portrait', 'B&W Filter', 'Cartoon Filter'],
        [
            ['Professional headshots', '✓ On', '✓ On', '✓ On'],
            ['Product photos', '✗ Off', '✓ On', '✓ On'],
            ['Artistic portraits', '✓ On', '✗ Off', '✓ On'],
            ['Character illustrations', '✗ Off', '✗ Off', '✗ Off']
        ]
    )
    
    doc.add_page_break()
    
    # Section 8: Output Files
    add_heading(doc, '8. Output Files', 1)
    
    add_heading(doc, 'Excel Output Structure', 2)
    p = doc.add_paragraph()
    p.add_run('The output Excel file contains:\n\n')
    p.add_run('1. All original data from the input file\n')
    p.add_run('2. New image columns based on your mappings\n')
    p.add_run('3. Images embedded directly in cells')
    
    add_heading(doc, 'Image Properties', 3)
    add_table(doc,
        ['Property', 'Value'],
        [
            ['Scale', '20% of original size'],
            ['Position', 'Anchored to cell'],
            ['Row Height', '120 pixels (auto-set)'],
            ['Column Width', '22 characters (auto-set)']
        ]
    )
    
    add_heading(doc, 'Log Files', 2)
    doc.add_paragraph('Normal Log (_log.txt):')
    p = doc.add_paragraph()
    p.add_run('[START] 2025-01-15 10:30:00\n').font.size = Pt(9)
    p.add_run('[LOG] Theme: headshot portrait closeup face\n')
    p.add_run('[SEARCH] Tom Hanks\n')
    p.add_run('[URLS] (Bing Images) 24 found...\n')
    p.add_run('Time taken: 0h 5m 23s')
    
    doc.add_page_break()
    
    # Section 9: Tips & Best Practices
    add_heading(doc, '9. Tips & Best Practices', 1)
    
    add_heading(doc, 'For Best Image Results', 2)
    add_table(doc,
        ['Tip', 'Why It Helps'],
        [
            ['Use specific search terms', '"John Smith CEO Microsoft" works better than just "John Smith"'],
            ['Choose appropriate themes', 'Match theme to content type'],
            ['Enable all filters for headshots', 'Removes unwanted image types'],
            ['Start with 720p', 'Good balance of quality and speed'],
            ['Use PNG format', 'Best quality, no compression artifacts']
        ]
    )
    
    add_heading(doc, 'For Faster Processing', 2)
    add_table(doc,
        ['Tip', 'Impact'],
        [
            ['Use Bing Images', 'Fastest and most reliable'],
            ['Lower resolution', 'Smaller downloads = faster'],
            ['Stable internet', 'Avoids timeout retries'],
            ['Close other browsers', 'More resources for Canvex']
        ]
    )
    
    add_heading(doc, 'For Large Files (1000+ rows)', 2)
    p = doc.add_paragraph()
    p.add_run('1. Process in batches — Split into smaller files\n')
    p.add_run('2. Use lower resolution — 480p is sufficient for previews\n')
    p.add_run('3. Choose JPG format — Smaller file sizes\n')
    p.add_run('4. Monitor progress — Cancel if stuck')
    
    doc.add_page_break()
    
    # Section 10: Troubleshooting
    add_heading(doc, '10. Troubleshooting', 1)
    
    add_heading(doc, 'Common Issues', 2)
    
    add_heading(doc, '❌ "No images found"', 3)
    p = doc.add_paragraph()
    p.add_run('Causes: ').bold = True
    p.add_run('Search term too vague, name misspelled, person not well-known\n\n')
    p.add_run('Solutions:\n').bold = True
    p.add_run('• Make search terms more specific\n')
    p.add_run('• Try a different search engine\n')
    p.add_run('• Simplify the theme\n')
    p.add_run('• Check spelling in Excel')
    
    add_heading(doc, '❌ "Wrong images appearing"', 3)
    p = doc.add_paragraph()
    p.add_run('Causes: ').bold = True
    p.add_run('Common name, theme not matching content\n\n')
    p.add_run('Solutions:\n').bold = True
    p.add_run('• Add context: "John Smith actor" or "John Smith CEO"\n')
    p.add_run('• Try a different theme\n')
    p.add_run('• Use custom theme with specific keywords')
    
    add_heading(doc, '❌ "Processing is very slow"', 3)
    p = doc.add_paragraph()
    p.add_run('Causes: ').bold = True
    p.add_run('High resolution selected, slow internet, many rows\n\n')
    p.add_run('Solutions:\n').bold = True
    p.add_run('• Lower resolution to 480p or 720p\n')
    p.add_run('• Check internet speed\n')
    p.add_run('• Process in smaller batches')
    
    add_heading(doc, 'Error Messages', 2)
    add_table(doc,
        ['Error', 'Meaning', 'Solution'],
        [
            ['Chrome not found', 'ChromeDriver issue', 'Install/update Chrome browser'],
            ['Connection timeout', 'Network issue', 'Check internet, try again'],
            ['Permission denied', 'File locked', 'Close the Excel file'],
            ['Out of memory', 'Too many images', 'Process smaller batches']
        ]
    )
    
    doc.add_page_break()
    
    # Section 11: FAQ
    add_heading(doc, '11. FAQ', 1)
    
    add_heading(doc, 'General Questions', 2)
    
    p = doc.add_paragraph()
    p.add_run('Q: What file formats are supported?\n').bold = True
    p.add_run('A: Input must be .xlsx (Excel 2007+). Output is always .xlsx.\n\n')
    
    p.add_run('Q: Can I process multiple Excel files at once?\n').bold = True
    p.add_run('A: No, process one file at a time.\n\n')
    
    p.add_run('Q: Are my images saved locally?\n').bold = True
    p.add_run('A: Yes, images are embedded directly in the output Excel file.\n\n')
    
    p.add_run('Q: Does Canvex work offline?\n').bold = True
    p.add_run('A: No, internet connection is required for image searches.')
    
    add_heading(doc, 'Technical Questions', 2)
    
    p = doc.add_paragraph()
    p.add_run('Q: Why does Canvex need Chrome?\n').bold = True
    p.add_run('A: Canvex uses Selenium with Chrome to scrape image search results.\n\n')
    
    p.add_run('Q: Where are settings saved?\n').bold = True
    p.add_run('A: macOS: ~/Library/Application Support/Canvex/\n')
    p.add_run('    Windows: %APPDATA%/Canvex/\n\n')
    
    p.add_run('Q: How long does processing take?\n').bold = True
    p.add_run('A: Typically ~2-5 seconds per row at 720p resolution.')
    
    doc.add_page_break()
    
    # Contact
    add_heading(doc, 'Contact & Support', 1)
    
    p = doc.add_paragraph()
    p.add_run('Publisher: ').bold = True
    p.add_run('Kunal Pagariya\n')
    p.add_run('Email: ').bold = True
    p.add_run('kunal.pagariya@outlook.com\n')
    p.add_run('Version: ').bold = True
    p.add_run('1.0\n\n')
    p.add_run('© 2025 Kunal Pagariya')
    
    # Final note
    doc.add_paragraph()
    final = doc.add_paragraph('Thank you for using Canvex!')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.runs[0].italic = True
    
    return doc


def create_technical_docs():
    """Create the Technical Documentation Word document"""
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Canvex', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(36)
    
    subtitle = doc.add_paragraph('Functional & Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = GRAY_TEXT
    
    version = doc.add_paragraph('Version 1.0 | Last Updated: December 2025')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.color.rgb = GRAY_TEXT
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    
    toc_part1 = doc.add_paragraph()
    toc_part1.add_run('Part I: Functional Specification\n').bold = True
    toc_part1.add_run('1. Product Overview\n')
    toc_part1.add_run('2. User Stories & Use Cases\n')
    toc_part1.add_run('3. Functional Requirements\n')
    toc_part1.add_run('4. User Interface Specification\n')
    toc_part1.add_run('5. Workflow & Process Flows\n')
    
    toc_part2 = doc.add_paragraph()
    toc_part2.add_run('Part II: Technical Specification\n').bold = True
    toc_part2.add_run('6. Architecture Overview\n')
    toc_part2.add_run('7. Module Documentation\n')
    toc_part2.add_run('8. Data Flow & Processing\n')
    toc_part2.add_run('9. External Dependencies\n')
    toc_part2.add_run('10. Configuration & Settings\n')
    toc_part2.add_run('11. Error Handling & Logging\n')
    toc_part2.add_run('12. Performance Optimization\n')
    toc_part2.add_run('13. Security Considerations\n')
    toc_part2.add_run('14. Deployment & Packaging')
    
    doc.add_page_break()
    
    # PART I Header
    part1 = doc.add_heading('Part I: Functional Specification', 1)
    part1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in part1.runs:
        run.font.size = Pt(28)
    
    doc.add_page_break()
    
    # Section 1: Product Overview
    add_heading(doc, '1. Product Overview', 1)
    
    add_heading(doc, '1.1 Purpose', 2)
    doc.add_paragraph(
        'Canvex is a desktop application designed to automate the process of searching for images '
        'on the web and inserting them into Excel spreadsheets based on text data in specified columns.'
    )
    
    add_heading(doc, '1.2 Problem Statement', 2)
    doc.add_paragraph('Manually searching for images and inserting them into Excel files is:')
    p = doc.add_paragraph()
    p.add_run('• Time-consuming: ').bold = True
    p.add_run('Each image requires multiple steps\n')
    p.add_run('• Error-prone: ').bold = True
    p.add_run('Easy to mix up images with wrong entries\n')
    p.add_run('• Tedious: ').bold = True
    p.add_run('Repetitive for large datasets')
    
    add_heading(doc, '1.3 Solution', 2)
    doc.add_paragraph('Canvex automates this workflow:')
    p = doc.add_paragraph()
    p.add_run('Excel Data → Image Search → Download → Filter → Resize → Insert → Export')
    p.runs[0].font.size = Pt(11)
    p.runs[0].bold = True
    
    add_heading(doc, '1.4 Target Users', 2)
    add_table(doc,
        ['User Type', 'Description', 'Primary Use Case'],
        [
            ['HR Professionals', 'Create employee directories', 'Headshot directories'],
            ['Content Creators', 'Build media catalogs', 'Actor/character sheets'],
            ['Marketers', 'Product catalogs', 'Product image galleries'],
            ['Researchers', 'Data visualization', 'Image-rich datasets']
        ]
    )
    
    add_heading(doc, '1.5 Key Value Propositions', 2)
    p = doc.add_paragraph()
    p.add_run('1. Automation ').bold = True
    p.add_run('— Reduces hours of manual work to minutes\n')
    p.add_run('2. Intelligence ').bold = True
    p.add_run('— Smart filtering removes unwanted images\n')
    p.add_run('3. Flexibility ').bold = True
    p.add_run('— Multiple search engines, themes, and output options\n')
    p.add_run('4. Reliability ').bold = True
    p.add_run('— Checkpoint saving, error recovery, and logging')
    
    doc.add_page_break()
    
    # Section 2: User Stories
    add_heading(doc, '2. User Stories & Use Cases', 1)
    
    add_heading(doc, '2.1 User Stories', 2)
    
    add_heading(doc, 'US-001: Basic Image Insertion', 3)
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('user with an Excel file containing names,\n')
    p.add_run('I want to ').italic = True
    p.add_run('automatically find and insert headshot images,\n')
    p.add_run('So that ').italic = True
    p.add_run('I can create a visual directory without manual searching.')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Acceptance Criteria:\n').bold = True
    p.add_run('• User can load an Excel file\n')
    p.add_run('• User can specify which column contains search terms\n')
    p.add_run('• Application searches for images and inserts them\n')
    p.add_run('• Output Excel contains embedded images')
    
    add_heading(doc, 'US-002: Custom Search Configuration', 3)
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('user processing different types of content,\n')
    p.add_run('I want to ').italic = True
    p.add_run('customize search parameters (theme, resolution, format),\n')
    p.add_run('So that ').italic = True
    p.add_run('I get appropriate images for my specific use case.')
    
    add_heading(doc, 'US-003: Multiple Column Processing', 3)
    p = doc.add_paragraph()
    p.add_run('As a ').italic = True
    p.add_run('user with multiple columns needing images,\n')
    p.add_run('I want to ').italic = True
    p.add_run('create multiple mappings in one run,\n')
    p.add_run('So that ').italic = True
    p.add_run("I don't have to process the same file multiple times.")
    
    doc.add_page_break()
    
    # Section 3: Functional Requirements
    add_heading(doc, '3. Functional Requirements', 1)
    
    add_heading(doc, '3.1 File Operations', 2)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-001', 'System shall accept .xlsx files as input', 'Must'],
            ['FR-002', 'System shall support drag-and-drop file loading', 'Must'],
            ['FR-003', 'System shall detect and list multiple sheets', 'Must'],
            ['FR-004', 'System shall export output as .xlsx with embedded images', 'Must'],
            ['FR-005', 'System shall create processing log files', 'Must'],
            ['FR-006', 'System shall maintain list of recently opened files', 'Should']
        ]
    )
    
    add_heading(doc, '3.2 Image Search', 2)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-010', 'System shall search Bing Images', 'Must'],
            ['FR-011', 'System shall search Google Images', 'Must'],
            ['FR-012', 'System shall search DuckDuckGo Images', 'Should'],
            ['FR-013', 'System shall append theme keywords to search queries', 'Must'],
            ['FR-014', 'System shall filter out stock photo websites', 'Must'],
            ['FR-015', 'System shall retry failed searches with fallback engine', 'Should']
        ]
    )
    
    add_heading(doc, '3.3 Image Processing', 2)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['FR-020', 'System shall download images in parallel', 'Must'],
            ['FR-021', 'System shall resize images to target resolution', 'Must'],
            ['FR-022', 'System shall filter low-quality images', 'Must'],
            ['FR-023', 'System shall filter black/white images (optional)', 'Should'],
            ['FR-024', 'System shall filter cartoon/graphic images (optional)', 'Should'],
            ['FR-025', 'System shall prefer portrait-oriented images (optional)', 'Should']
        ]
    )
    
    add_heading(doc, '3.4 Non-Functional Requirements', 2)
    add_table(doc,
        ['ID', 'Requirement', 'Priority'],
        [
            ['NFR-001', 'System shall process at least 2-5 rows per second', 'Should'],
            ['NFR-002', 'System shall support files with 10,000+ rows', 'Should'],
            ['NFR-003', 'System shall run on macOS 10.14+ and Windows 10+', 'Must'],
            ['NFR-004', 'System shall provide clear error messages', 'Must'],
            ['NFR-005', 'System shall checkpoint progress on cancellation', 'Must']
        ]
    )
    
    doc.add_page_break()
    
    # Section 4: UI Specification
    add_heading(doc, '4. User Interface Specification', 1)
    
    add_heading(doc, '4.1 Main Window Properties', 2)
    add_table(doc,
        ['Property', 'Value'],
        [
            ['Title', 'Canvex'],
            ['Default Size', '900 × 700 pixels'],
            ['Minimum Size', '700 × 600 pixels'],
            ['Resizable', 'Yes'],
            ['Accept Drops', 'Yes (.xlsx files)']
        ]
    )
    
    add_heading(doc, '4.2 Theme Colors', 2)
    add_table(doc,
        ['Element', 'Dark Theme', 'Light Theme'],
        [
            ['Background', '#1e1e1e', '#f5f5f7'],
            ['Card Background', '#2d2d2d', '#ffffff'],
            ['Text', '#ffffff', '#1d1d1f'],
            ['Accent', '#0a84ff', '#0a84ff'],
            ['Border', '#404040', '#d2d2d7'],
            ['Error', '#ff3b30', '#ff3b30'],
            ['Success', '#34c759', '#34c759']
        ]
    )
    
    doc.add_page_break()
    
    # Section 5: Workflow
    add_heading(doc, '5. Workflow & Process Flows', 1)
    
    add_heading(doc, '5.1 High-Level Workflow', 2)
    doc.add_paragraph('The application follows this main workflow:')
    p = doc.add_paragraph()
    p.add_run('STARTUP').bold = True
    p.add_run(' → ')
    p.add_run('LOAD FILE').bold = True
    p.add_run(' → ')
    p.add_run('CONFIGURE').bold = True
    p.add_run(' → ')
    p.add_run('PROCESS').bold = True
    p.add_run(' → ')
    p.add_run('COMPLETE').bold = True
    
    add_heading(doc, '5.2 Processing Steps', 2)
    p = doc.add_paragraph()
    p.add_run('For each row in Excel:\n').bold = True
    p.add_run('1. Write text data to output workbook\n')
    p.add_run('2. For each mapping:\n')
    p.add_run('   a. Search images using selected browser\n')
    p.add_run('   b. Download candidates in parallel (max 8)\n')
    p.add_run('   c. Apply filters (brightness, color, portrait)\n')
    p.add_run('   d. Select best image\n')
    p.add_run('   e. Save temp file and insert into Excel\n')
    p.add_run('3. Update progress bar\n')
    p.add_run('4. Check for cancellation request')
    
    doc.add_page_break()
    
    # PART II Header
    part2 = doc.add_heading('Part II: Technical Specification', 1)
    part2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in part2.runs:
        run.font.size = Pt(28)
    
    doc.add_page_break()
    
    # Section 6: Architecture
    add_heading(doc, '6. Architecture Overview', 1)
    
    add_heading(doc, '6.1 Application Layers', 2)
    p = doc.add_paragraph()
    p.add_run('Presentation Layer\n').bold = True
    p.add_run('• Main Window (PyQt5 QWidget)\n')
    p.add_run('• Dialogs (Settings, Help, About)\n')
    p.add_run('• Theme Management\n\n')
    
    p.add_run('Business Logic Layer\n').bold = True
    p.add_run('• WorkerUltra (QThread)\n')
    p.add_run('• Settings Manager\n')
    p.add_run('• Column Mapping Handler\n\n')
    
    p.add_run('Data/Service Layer\n').bold = True
    p.add_run('• Image Search (Selenium)\n')
    p.add_run('• Image Processing (Pillow)\n')
    p.add_run('• Excel I/O (pandas/xlsxwriter)')
    
    add_heading(doc, '6.2 Threading Model', 2)
    p = doc.add_paragraph()
    p.add_run('Main Thread (GUI):\n').bold = True
    p.add_run('• PyQt5 event loop\n')
    p.add_run('• User interaction handling\n')
    p.add_run('• UI updates via signals\n\n')
    
    p.add_run('Worker Thread (WorkerUltra):\n').bold = True
    p.add_run('• Heavy I/O operations\n')
    p.add_run('• Selenium browser control\n')
    p.add_run('• Network requests\n')
    p.add_run('• File writes\n\n')
    
    p.add_run('Thread Pool:\n').bold = True
    p.add_run('• Parallel image downloads\n')
    p.add_run('• Max workers: min(20, CPU_count * 2)')
    
    doc.add_page_break()
    
    # Section 7: Module Documentation
    add_heading(doc, '7. Module Documentation', 1)
    
    add_heading(doc, '7.1 Main Application Class', 2)
    p = doc.add_paragraph()
    p.add_run('CanvaImageExcelCreator(QWidget)\n').bold = True
    p.add_run('Purpose: Main GUI window and application controller\n\n')
    p.add_run('Key Attributes:\n').bold = True
    p.add_run('• excel_path: str - Path to loaded Excel file\n')
    p.add_run('• columns: list[str] - Column names from Excel\n')
    p.add_run('• worker: WorkerUltra - Background processing thread\n')
    p.add_run('• session_running: bool - Processing state flag\n')
    p.add_run('• settings_path: str - Path to settings JSON file')
    
    add_heading(doc, '7.2 Worker Thread Class', 2)
    p = doc.add_paragraph()
    p.add_run('WorkerUltra(QThread)\n').bold = True
    p.add_run('Purpose: Background thread for heavy processing operations\n\n')
    p.add_run('Signals:\n').bold = True
    
    add_table(doc,
        ['Signal', 'Type', 'Purpose'],
        [
            ['sig_overall', 'pyqtSignal(int)', 'Overall progress 0-100'],
            ['sig_step', 'pyqtSignal(int)', 'Per-item progress 0-100'],
            ['sig_log', 'pyqtSignal(str)', 'Log message string'],
            ['sig_done', 'pyqtSignal(str)', 'Success with output path'],
            ['sig_error', 'pyqtSignal(str)', 'Error message']
        ]
    )
    
    add_heading(doc, '7.3 Image Search Functions', 2)
    add_table(doc,
        ['Function', 'Parameters', 'Returns'],
        [
            ['bing_urls()', 'driver, term, theme, limit', 'list[str] - Image URLs'],
            ['google_urls()', 'driver, term, theme, limit', 'list[str] - Image URLs'],
            ['ddg_urls()', 'driver, term, theme, limit', 'list[str] - Image URLs'],
            ['fetch_image_urls()', 'driver, term, theme, browser, limit', 'list[str] - URLs from selected engine']
        ]
    )
    
    add_heading(doc, '7.4 Image Processing Functions', 2)
    add_table(doc,
        ['Function', 'Purpose'],
        [
            ['dl_resize(url, target)', 'Download and resize image with quality filtering'],
            ['is_valid_image_url(url)', 'Validate URL for processing'],
            ['create_driver()', 'Create configured Chrome WebDriver']
        ]
    )
    
    doc.add_page_break()
    
    # Section 8: Data Flow
    add_heading(doc, '8. Data Flow & Processing', 1)
    
    add_heading(doc, '8.1 Input Data Flow', 2)
    p = doc.add_paragraph()
    p.add_run('Excel File (.xlsx)\n').bold = True
    p.add_run('    ↓ pandas.read_excel()\n')
    p.add_run('DataFrame with columns\n').bold = True
    p.add_run('    ↓ User Configuration\n')
    p.add_run('WorkerUltra.__init__()\n').bold = True
    
    add_heading(doc, '8.2 Excel Output Structure', 2)
    add_table(doc,
        ['Property', 'Value'],
        [
            ['Sheet Name', 'Based on output filename (max 31 chars)'],
            ['Image Scale', '20% (x_scale: 0.20, y_scale: 0.20)'],
            ['Object Position', '1 (move with cells)'],
            ['Column Width', '22 characters (for image columns)'],
            ['Row Height', '120 pixels (default)']
        ]
    )
    
    doc.add_page_break()
    
    # Section 9: Dependencies
    add_heading(doc, '9. External Dependencies', 1)
    
    add_heading(doc, '9.1 Python Package Dependencies', 2)
    add_table(doc,
        ['Package', 'Version', 'Purpose'],
        [
            ['PyQt5', '≥5.15', 'GUI framework'],
            ['pandas', '≥1.3', 'Excel reading'],
            ['xlsxwriter', '≥3.0', 'Excel writing with images'],
            ['Pillow', '≥9.0', 'Image processing'],
            ['selenium', '≥4.0', 'Web scraping'],
            ['webdriver-manager', '≥3.8', 'ChromeDriver management'],
            ['requests', '≥2.28', 'HTTP requests'],
            ['qtawesome', '≥1.0', '(Optional) Font Awesome icons']
        ]
    )
    
    add_heading(doc, '9.2 System Dependencies', 2)
    add_table(doc,
        ['Dependency', 'Purpose', 'Required'],
        [
            ['Chrome Browser', 'Selenium web scraping', 'Yes'],
            ['ChromeDriver', 'Chrome automation', 'Auto-installed']
        ]
    )
    
    add_heading(doc, '9.3 Blocked Domains (BAD_SITES)', 2)
    doc.add_paragraph('Images from these domains are filtered out:')
    p = doc.add_paragraph()
    p.add_run('shutterstock, alamy, getty, adobe, dreamstime, depositphotos, ')
    p.add_run('123rf, bigstock, vectorstock, istock')
    
    doc.add_page_break()
    
    # Section 10: Configuration
    add_heading(doc, '10. Configuration & Settings', 1)
    
    add_heading(doc, '10.1 Settings File Location', 2)
    add_table(doc,
        ['Platform', 'Path'],
        [
            ['macOS (bundled)', '~/Library/Application Support/Canvex/canva_last_settings.json'],
            ['Windows (bundled)', '%APPDATA%/Canvex/canva_last_settings.json'],
            ['Development', './canva_last_settings.json']
        ]
    )
    
    add_heading(doc, '10.2 Settings Schema', 2)
    add_table(doc,
        ['Property', 'Type', 'Default'],
        [
            ['theme', 'string', 'headshot portrait closeup face'],
            ['custom_theme', 'string', '""'],
            ['resolution', 'string', '720p'],
            ['format', 'string', 'PNG'],
            ['jpg_quality', 'integer', '90'],
            ['browser', 'string', 'Bing Images'],
            ['filter_portrait', 'boolean', 'true'],
            ['filter_bw', 'boolean', 'true'],
            ['filter_cartoon', 'boolean', 'true']
        ]
    )
    
    doc.add_page_break()
    
    # Section 11: Error Handling
    add_heading(doc, '11. Error Handling & Logging', 1)
    
    add_heading(doc, '11.1 Error Handling Strategy', 2)
    
    p = doc.add_paragraph()
    p.add_run('Level 1: Image Download Errors\n').bold = True
    p.add_run('• Retry twice with 0.2s backoff\n')
    p.add_run('• On failure: Return None, try next URL\n\n')
    
    p.add_run('Level 2: Search Errors\n').bold = True
    p.add_run('• Retry page load once\n')
    p.add_run('• Google failure: Fallback to Bing\n\n')
    
    p.add_run('Level 3: Processing Errors\n').bold = True
    p.add_run('• Caught in try/except/finally\n')
    p.add_run('• Workbook always closed (checkpoint)\n')
    p.add_run('• Error log file written\n\n')
    
    p.add_run('Level 4: Application Errors\n').bold = True
    p.add_run('• safe_exit() prevents crash on force quit\n')
    p.add_run('• KeyboardInterrupt handled gracefully')
    
    add_heading(doc, '11.2 Log File Structure', 2)
    doc.add_paragraph('Normal Log (_log.txt):')
    p = doc.add_paragraph()
    p.add_run('[START] 2025-01-15 10:30:00\n').font.size = Pt(9)
    p.add_run('[LOG] Theme: headshot portrait closeup face\n')
    p.add_run('[SEARCH] Tom Hanks\n')
    p.add_run('[URLS] (Bing Images) 24 found...\n')
    p.add_run('Time taken: 0h 5m 23s')
    
    doc.add_page_break()
    
    # Section 12: Performance
    add_heading(doc, '12. Performance Optimization', 1)
    
    add_heading(doc, '12.1 Parallel Processing', 2)
    p = doc.add_paragraph()
    p.add_run('Thread Pool Configuration:\n').bold = True
    p.add_run('cpus = os.cpu_count() or 4\n')
    p.add_run('maxw = min(20, max(6, cpus * 2))  # 6-20 workers')
    
    add_heading(doc, '12.2 HTTP Connection Pooling', 2)
    p = doc.add_paragraph()
    p.add_run('• Global session with connection pooling\n')
    p.add_run('• Pool connections: 100\n')
    p.add_run('• Pool maxsize: 100\n')
    p.add_run('• Retry: 2 attempts with 0.2s backoff')
    
    add_heading(doc, '12.3 Image Caching', 2)
    p = doc.add_paragraph()
    p.add_run('• In-memory cache for downloaded images\n')
    p.add_run('• Cache limit: 1MB per image\n')
    p.add_run('• Prevents re-downloading duplicate URLs')
    
    add_heading(doc, '12.4 Selenium Optimization', 2)
    p = doc.add_paragraph()
    p.add_run('Chrome Options:\n').bold = True
    p.add_run('• --disable-gpu\n')
    p.add_run('• --no-sandbox\n')
    p.add_run('• --disable-extensions\n')
    p.add_run('• --incognito\n')
    p.add_run('• Page load timeout: 12 seconds')
    
    doc.add_page_break()
    
    # Section 13: Security
    add_heading(doc, '13. Security Considerations', 1)
    
    add_heading(doc, '13.1 Network Security', 2)
    add_table(doc,
        ['Risk', 'Mitigation'],
        [
            ['Malicious image URLs', 'URL validation before download'],
            ['Timeout attacks', '7-second timeout on requests'],
            ['SSL verification', 'Default SSL verification enabled']
        ]
    )
    
    add_heading(doc, '13.2 File System Security', 2)
    add_table(doc,
        ['Risk', 'Mitigation'],
        [
            ['Path traversal', 'Uses os.path.join, no user paths in temp names'],
            ['Temp file exposure', 'Temp files in system temp dir, cleaned on success'],
            ['Settings tampering', 'JSON parsing with exception handling']
        ]
    )
    
    doc.add_page_break()
    
    # Section 14: Deployment
    add_heading(doc, '14. Deployment & Packaging', 1)
    
    add_heading(doc, '14.1 PyInstaller Configuration', 2)
    doc.add_paragraph('Build commands:')
    p = doc.add_paragraph()
    p.add_run('macOS: ').bold = True
    p.add_run('pyinstaller Canvex.spec --noconfirm\n')
    p.add_run('Output: dist/Canvex.app\n\n')
    p.add_run('Windows: ').bold = True
    p.add_run('pyinstaller Canvex.spec --noconfirm\n')
    p.add_run('Output: dist/Canvex/Canvex.exe')
    
    add_heading(doc, '14.2 Required Assets', 2)
    add_table(doc,
        ['File', 'Purpose', 'Required'],
        [
            ['app_icon.ico', 'Windows icon', 'Yes'],
            ['app_icon.icns', 'macOS icon', 'Yes (for .app)'],
            ['splash.png', 'Splash screen', 'Optional'],
            ['logo.svg', 'Alternative logo', 'Optional']
        ]
    )
    
    add_heading(doc, '14.3 Distribution Checklist', 2)
    p = doc.add_paragraph()
    p.add_run('☐ Test on clean macOS installation\n')
    p.add_run('☐ Test on clean Windows installation\n')
    p.add_run('☐ Verify Chrome/ChromeDriver compatibility\n')
    p.add_run('☐ Check code signing (macOS notarization)\n')
    p.add_run('☐ Verify file associations work\n')
    p.add_run('☐ Test with various Excel files\n')
    p.add_run('☐ Verify settings persistence\n')
    p.add_run('☐ Check temp file cleanup')
    
    doc.add_page_break()
    
    # Appendices
    add_heading(doc, 'Appendix A: File Structure', 1)
    p = doc.add_paragraph()
    p.add_run('Canvex/\n').bold = True
    p.add_run('├── Canvex.py               # Main application source\n')
    p.add_run('├── Canvex.spec             # PyInstaller spec file\n')
    p.add_run('├── canva_last_settings.json # User settings\n')
    p.add_run('├── app_icon.ico            # Windows icon\n')
    p.add_run('├── app_icon.icns           # macOS icon\n')
    p.add_run('├── splash.png              # Splash screen image\n')
    p.add_run('├── docs/\n')
    p.add_run('│   ├── USER_GUIDE.md\n')
    p.add_run('│   └── TECHNICAL_DOCS.md\n')
    p.add_run('└── build/                  # PyInstaller output')
    
    add_heading(doc, 'Appendix B: Supported Formats', 1)
    
    add_heading(doc, 'Input Formats', 2)
    add_table(doc,
        ['Format', 'Extension', 'Support'],
        [
            ['Excel 2007+', '.xlsx', '✓ Full'],
            ['Excel 97-2003', '.xls', '✗ Not supported'],
            ['CSV', '.csv', '✗ Not supported']
        ]
    )
    
    add_heading(doc, 'Output Image Formats', 2)
    add_table(doc,
        ['Format', 'Quality', 'Transparency', 'Notes'],
        [
            ['PNG', 'Lossless', '✓ Yes', 'Best quality, larger files'],
            ['JPEG', '60-100', '✗ No', 'Configurable quality'],
            ['WebP', '95', '✓ Yes', 'Modern format, best compression']
        ]
    )
    
    doc.add_page_break()
    
    # Final page
    final = doc.add_paragraph('End of Technical Documentation')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.runs[0].font.size = Pt(14)
    final.runs[0].italic = True
    
    version = doc.add_paragraph('Version 1.0 | © 2025 Kunal Pagariya')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.runs[0].font.color.rgb = GRAY_TEXT
    
    return doc


def main():
    """Generate both Word documents"""
    output_dir = os.path.dirname(os.path.abspath(__file__))
    
    print("Generating Canvex documentation...")
    
    # Create User Guide
    print("  Creating USER_GUIDE.docx...")
    user_guide = create_user_guide()
    user_guide_path = os.path.join(output_dir, 'Canvex_User_Guide.docx')
    user_guide.save(user_guide_path)
    print(f"  ✓ Saved: {user_guide_path}")
    
    # Create Technical Docs
    print("  Creating TECHNICAL_DOCS.docx...")
    tech_docs = create_technical_docs()
    tech_docs_path = os.path.join(output_dir, 'Canvex_Technical_Documentation.docx')
    tech_docs.save(tech_docs_path)
    print(f"  ✓ Saved: {tech_docs_path}")
    
    print("\n✅ Documentation generated successfully!")
    print(f"\nFiles created in: {output_dir}")


if __name__ == "__main__":
    main()
